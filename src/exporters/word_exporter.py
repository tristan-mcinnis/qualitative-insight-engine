from pathlib import Path
from datetime import datetime
import pandas as pd
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from src.config.config_loader import AppConfig
from src.utils.logger import Logger

class ColorSchemeManager:
    TOPIC_COLORS = [
        RGBColor(0x1f, 0x77, 0xb4),  # Blue
        RGBColor(0xff, 0x7f, 0x0e),  # Orange  
        RGBColor(0x2c, 0xa0, 0x2c),  # Green
        RGBColor(0xd6, 0x27, 0x28),  # Red
        RGBColor(0x94, 0x67, 0xbd),  # Purple
        RGBColor(0x8c, 0x56, 0x4b),  # Brown
        RGBColor(0xe3, 0x77, 0xc2),  # Pink
        RGBColor(0x7f, 0x7f, 0x7f),  # Gray
        RGBColor(0xbc, 0xbd, 0x22),  # Olive
        RGBColor(0x17, 0xbe, 0xcf),  # Cyan
        RGBColor(0xae, 0xc7, 0xe8),  # Light Blue
        RGBColor(0xff, 0xbb, 0x78),  # Light Orange
        RGBColor(0x98, 0xdf, 0x8a),  # Light Green
        RGBColor(0xff, 0x98, 0x96),  # Light Red
        RGBColor(0xc5, 0xb0, 0xd5),  # Light Purple
    ]
    
    HIGHLIGHT_COLORS = [
        WD_COLOR_INDEX.TURQUOISE,
        WD_COLOR_INDEX.BRIGHT_GREEN,
        WD_COLOR_INDEX.PINK,
        WD_COLOR_INDEX.BLUE,
        WD_COLOR_INDEX.RED,
        WD_COLOR_INDEX.DARK_YELLOW,
        WD_COLOR_INDEX.TEAL,
        WD_COLOR_INDEX.GRAY_25,
        WD_COLOR_INDEX.YELLOW,
        WD_COLOR_INDEX.GREEN,
        WD_COLOR_INDEX.VIOLET,
        WD_COLOR_INDEX.DARK_BLUE,
        WD_COLOR_INDEX.DARK_RED,
        WD_COLOR_INDEX.WHITE,
        WD_COLOR_INDEX.GRAY_50
    ]
    
    def __init__(self, config: AppConfig, logger: Logger = None):
        self.config = config
        self.logger = logger or Logger()
        self.topic_to_color = {}
        self.topic_to_highlight = {}
        self.color_index = 0
    
    def get_color_for_topic(self, broad_topic: str, sub_topic: str = None) -> RGBColor:
        key = broad_topic.strip().lower()
        
        if key not in self.topic_to_color:
            if self.color_index >= self.config.word.max_topics_for_colors:
                self.logger.warning(
                    f"Topic {self.color_index + 1}: '{broad_topic}' - Color palette cycling"
                )
            
            color_idx = self.color_index % len(self.TOPIC_COLORS)
            highlight_idx = self.color_index % len(self.HIGHLIGHT_COLORS)
            
            self.topic_to_color[key] = self.TOPIC_COLORS[color_idx]
            self.topic_to_highlight[key] = self.HIGHLIGHT_COLORS[highlight_idx]
            self.color_index += 1
        
        return self.topic_to_color[key]
    
    def get_highlight_for_topic(self, broad_topic: str, sub_topic: str = None):
        key = broad_topic.strip().lower()
        self.get_color_for_topic(broad_topic, sub_topic)
        return self.topic_to_highlight.get(key, WD_COLOR_INDEX.YELLOW)

class WordExporter:
    def __init__(self, config: AppConfig, logger: Logger):
        self.config = config
        self.logger = logger
        self.color_manager = ColorSchemeManager(config, logger)
    
    def create_analysis_report(self, master_df: pd.DataFrame, strategic_df: pd.DataFrame,
                             output_path: Path, project_name: str = "Qualitative Analysis") -> None:
        if not self.config.word.enabled:
            return
        
        self.logger.info("Generating Word analysis report...")
        
        doc = docx.Document()
        
        self._add_title_page(doc, project_name)
        self._add_executive_summary(doc, strategic_df)
        self._add_color_legend(doc, master_df)
        self._add_coded_verbatims(doc, master_df)
        self._add_strategic_insights(doc, strategic_df)
        
        doc.save(str(output_path))
        self.logger.info(f"Word analysis report saved: {output_path.name}")
    
    def create_executive_summary(self, strategic_df: pd.DataFrame, output_path: Path,
                               project_name: str = "Qualitative Analysis") -> None:
        if not self.config.word.enabled:
            return
        
        self.logger.info("Generating executive summary...")
        
        doc = docx.Document()
        
        title = doc.add_heading(f"{project_name} - Executive Summary", 0)
        doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        if not strategic_df.empty:
            doc.add_heading("Key Strategic Insights", level=1)
            
            for broad_topic, group in strategic_df.groupby('Broad Topic'):
                if broad_topic and broad_topic != 'No Topic Assigned':
                    doc.add_heading(broad_topic, level=2)
                    
                    insights = group['Key Insights'].dropna().unique()
                    for insight in insights[:3]:
                        if insight:
                            p = doc.add_paragraph()
                            p.add_run("• ").bold = True
                            p.add_run(insight)
                    
                    doc.add_paragraph()
        
        doc.save(str(output_path))
        self.logger.info(f"Executive summary saved: {output_path.name}")
    
    def create_qda_workbook(self, master_df: pd.DataFrame, output_path: Path,
                          project_name: str = "Qualitative Analysis") -> None:
        if not self.config.word.enabled:
            return
        
        self.logger.info("Generating manual coding workbook...")
        
        doc = docx.Document()
        
        title = doc.add_heading(f"{project_name} - Manual Coding Workbook", 0)
        
        doc.add_heading("Instructions for Manual Coding", level=1)
        instructions = [
            "This workbook contains verbatims with AI-suggested topic coding.",
            "Review the highlighted text and topic assignments.",
            "Add your own codes using the blank coding tables provided.",
            "Use consistent color coding for your manual analysis.",
            "Document your coding decisions in the notes sections."
        ]
        
        for instruction in instructions:
            p = doc.add_paragraph()
            p.add_run("• ").bold = True
            p.add_run(instruction)
        
        doc.add_paragraph()
        
        self._add_coded_verbatims(doc, master_df, include_confidence=True)
        self._add_blank_coding_tables(doc)
        
        doc.save(str(output_path))
        self.logger.info(f"Manual coding workbook saved: {output_path.name}")
    
    def _add_title_page(self, doc, project_name: str) -> None:
        title = doc.add_heading(project_name, 0)
        title.alignment = 1
        
        subtitle = doc.add_heading("Qualitative Research Analysis Report", level=1)
        subtitle.alignment = 1
        
        doc.add_paragraph()
        
        meta_info = [
            f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}",
            f"Generated by: Qualitative Analysis Pipeline v{self.config.version}",
            f"Methodology: AI-Assisted Topic Analysis with Strategic Insights"
        ]
        
        for info in meta_info:
            p = doc.add_paragraph(info)
            p.alignment = 1
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, strategic_df: pd.DataFrame) -> None:
        doc.add_heading("Executive Summary", level=1)
        
        if strategic_df.empty:
            doc.add_paragraph("No strategic insights available for summary.")
            return
        
        doc.add_heading("Key Findings", level=2)
        
        total_topics = len(strategic_df['Broad Topic'].unique())
        total_insights = len(strategic_df)
        
        summary_stats = doc.add_paragraph()
        summary_stats.add_run(f"• Identified {total_topics} major themes\n")
        summary_stats.add_run(f"• Generated {total_insights} strategic insights\n")
        summary_stats.add_run(f"• Analysis includes color-coded verbatims and recommendations")
        
        doc.add_paragraph()
        
        doc.add_heading("Major Themes", level=2)
        
        for broad_topic, group in strategic_df.groupby('Broad Topic'):
            if broad_topic and broad_topic != 'No Topic Assigned':
                theme_para = doc.add_paragraph()
                theme_para.add_run(f"• {broad_topic}").bold = True
                
                first_insight = group['Key Insights'].dropna().iloc[0] if not group['Key Insights'].dropna().empty else ""
                if first_insight:
                    preview = first_insight[:200] + "..." if len(first_insight) > 200 else first_insight
                    theme_para.add_run(f": {preview}")
        
        doc.add_page_break()
    
    def _add_color_legend(self, doc, master_df: pd.DataFrame) -> None:
        doc.add_heading("Color Coding Legend", level=1)
        
        doc.add_paragraph("The following colors are used to highlight verbatims by topic:")
        doc.add_paragraph()
        
        topics = master_df['Emergent Broad Topic'].dropna().unique()
        topics = [t for t in topics if t != 'No Topic Assigned']
        
        for topic in sorted(topics):
            color = self.color_manager.get_color_for_topic(topic)
            
            p = doc.add_paragraph()
            colored_run = p.add_run(f"■ {topic}")
            colored_run.font.color.rgb = color
            colored_run.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph("Note: Sub-topics within each broad topic use the same color family.")
        doc.add_page_break()
    
    def _add_coded_verbatims(self, doc, master_df: pd.DataFrame, include_confidence: bool = False) -> None:
        doc.add_heading("Coded Verbatims by Topic", level=1)
        
        if master_df.empty:
            doc.add_paragraph("No verbatims available for coding.")
            return
        
        grouped = master_df.groupby('Emergent Broad Topic')
        
        for broad_topic, group in grouped:
            if broad_topic == 'No Topic Assigned':
                continue
                
            topic_header = doc.add_heading(broad_topic, level=2)
            color = self.color_manager.get_color_for_topic(broad_topic)
            topic_header.runs[0].font.color.rgb = color
            
            sub_grouped = group.groupby('Emergent Sub-Topic')
            
            for sub_topic, sub_group in sub_grouped:
                if sub_topic and sub_topic != 'No Sub-Topic Assigned':
                    sub_header = doc.add_heading(sub_topic, level=3)
                
                for _, row in sub_group.iterrows():
                    if pd.notna(row['Verbatim']):
                        p = doc.add_paragraph()
                        
                        speaker_run = p.add_run(f"[{row['Speaker']}]: ")
                        speaker_run.bold = True
                        
                        verbatim_run = p.add_run(row['Verbatim'])
                        highlight_color = self.color_manager.get_highlight_for_topic(broad_topic, sub_topic)
                        verbatim_run.font.highlight_color = highlight_color
                        
                        if pd.notna(row.get('Verbatim (English)')) and row.get('Verbatim (English)') != row['Verbatim']:
                            p.add_run("\n")
                            translation_run = p.add_run(f"Translation: {row['Verbatim (English)']}")
                            translation_run.italic = True
                        
                        if include_confidence and pd.notna(row.get('Confidence')):
                            p.add_run(f"\n[AI Confidence: {row['Confidence']}]")
                        
                        if pd.notna(row.get('Source File')):
                            source_run = p.add_run(f"\n(Source: {row['Source File']})")
                            source_run.font.size = Pt(9)
                        
                        doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_strategic_insights(self, doc, strategic_df: pd.DataFrame) -> None:
        doc.add_heading("Strategic Analysis & Recommendations", level=1)
        
        if strategic_df.empty:
            doc.add_paragraph("No strategic insights available.")
            return
        
        for broad_topic, group in strategic_df.groupby('Broad Topic'):
            if broad_topic == 'No Topic Assigned':
                continue
                
            topic_header = doc.add_heading(broad_topic, level=2)
            color = self.color_manager.get_color_for_topic(broad_topic)
            topic_header.runs[0].font.color.rgb = color
            
            insights = group['Key Insights'].dropna().unique()
            themes = group['Theme'].dropna().unique()
            takeaways = group['Takeaway'].dropna().unique()
            quotes = group['Supporting Quote'].dropna().unique()
            
            if len(insights) > 0:
                doc.add_heading("Key Insights", level=3)
                for insight in insights:
                    if insight:
                        p = doc.add_paragraph(insight)
                        p.style = 'List Bullet'
            
            if len(themes) > 0:
                doc.add_heading("Key Themes", level=3)
                for theme in themes:
                    if theme:
                        p = doc.add_paragraph(theme)
                        p.style = 'List Bullet'
            
            if len(takeaways) > 0:
                doc.add_heading("Strategic Recommendations", level=3)
                for takeaway in takeaways:
                    if takeaway:
                        p = doc.add_paragraph(takeaway)
                        p.style = 'List Bullet'
            
            if len(quotes) > 0:
                doc.add_heading("Supporting Evidence", level=3)
                for quote in quotes:
                    if quote:
                        p = doc.add_paragraph()
                        quote_run = p.add_run(f'"{quote}"')
                        quote_run.italic = True
            
            doc.add_paragraph()
    
    def _add_blank_coding_tables(self, doc) -> None:
        doc.add_heading("Manual Coding Workspace", level=1)
        
        doc.add_paragraph("Use the tables below for your manual coding analysis:")
        doc.add_paragraph()
        
        doc.add_heading("Theme Coding Table", level=2)
        table = doc.add_table(rows=10, cols=3)
        table.style = 'Table Grid'
        
        headers = ["Quote/Verbatim", "Code/Theme", "Notes"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        doc.add_heading("Analysis Notes", level=2)
        doc.add_paragraph("_" * 80)
        for i in range(15):
            doc.add_paragraph("_" * 80)