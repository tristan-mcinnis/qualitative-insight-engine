from pathlib import Path
from typing import List, Dict, Any
import docx
from src.config.config_loader import AppConfig
from src.utils.logger import Logger
from src.processors.text_processor import TextProcessor

class FileReader:
    ENCODINGS = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
    
    @classmethod
    def read_text_file(cls, file_path: Path) -> str:
        for encoding in cls.ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise UnicodeDecodeError(
            f"Could not decode file {file_path} with any common encoding"
        )
    
    @classmethod
    def read_docx_file(cls, file_path: Path) -> str:
        doc = docx.Document(file_path)
        return '\n'.join(para.text for para in doc.paragraphs)
    
    @classmethod
    def read_document(cls, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return cls.read_docx_file(file_path)
        elif suffix == '.txt':
            return cls.read_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

class VerbatimExtractor:
    def __init__(self, text_processor: TextProcessor, logger: Logger):
        self.text_processor = text_processor
        self.logger = logger
    
    def extract_from_file(self, file_path: Path) -> List[Dict[str, Any]]:
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return self._extract_from_docx(file_path)
        elif suffix == '.txt':
            return self._extract_from_txt(file_path)
        else:
            self.logger.warning(f"Unsupported file type: {suffix}")
            return []
    
    def _extract_from_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        verbatims = []
        doc = docx.Document(file_path)
        
        for para in doc.paragraphs:
            result = self.text_processor.extract_speaker_and_text(para.text)
            if result:
                speaker, text, timestamp = result
                verbatim = {
                    "speaker": speaker,
                    "text": text,
                    "source_file": file_path.name
                }
                if timestamp:
                    verbatim["timestamp"] = timestamp
                verbatims.append(verbatim)
        
        return verbatims
    
    def _extract_from_txt(self, file_path: Path) -> List[Dict[str, Any]]:
        verbatims = []
        content = FileReader.read_text_file(file_path)
        lines = content.split('\n')
        
        start_idx = self._find_content_start(lines)
        
        for line in lines[start_idx:]:
            result = self.text_processor.extract_speaker_and_text(line)
            if result:
                speaker, text, timestamp = result
                verbatim = {
                    "speaker": speaker,
                    "text": text,
                    "source_file": file_path.name
                }
                if timestamp:
                    verbatim["timestamp"] = timestamp
                verbatims.append(verbatim)
        
        return verbatims
    
    def _find_content_start(self, lines: List[str]) -> int:
        for i, line in enumerate(lines):
            if '=' * 10 in line and i > 0:
                return i + 1
        return 0